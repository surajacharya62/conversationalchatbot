import os
from typing import List
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.schema import Document
import shutil

#---- Document processing
class DocumentProcessor:
    def __init__(self, google_api_key: str):
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=google_api_key
        )
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,  
            chunk_overlap=50,  
            length_function=len,
        )
        self.vectorstore = None

    def setup_documents(self, uploaded_files):
        """Setup documents directly from Streamlit uploaded files - NO TEMP FILES"""
        try:
            print(f"🔧 Processing {len(uploaded_files)} files directly...")
            
            # Clear old documents first
            self.clear_vectorstore()
            
            documents = []
            
            for uploaded_file in uploaded_files:
                try:
                    print(f"📄 Processing: {uploaded_file.name}")
                    
                    # Get file content as text
                    content = ""
                    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
                    
                    # handling text files
                    if file_extension == '.txt':
                        
                        content = uploaded_file.getvalue().decode('utf-8')

                    # handling pdf files    
                    elif file_extension == '.pdf':
                        
                        try:
                            import PyPDF2
                            import io
                            
                            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.getvalue()))
                            text_parts = []
                            
                            for page in pdf_reader.pages:
                                page_text = page.extract_text()
                                if page_text and page_text.strip():
                                    text_parts.append(page_text.strip())
                            
                            content = "\n\n".join(text_parts)
                            
                        except Exception as pdf_error:
                            print(f"PDF extraction failed: {pdf_error}")
                            content = f"Error extracting text from {uploaded_file.name}"
                    
                    # handling docx file
                    elif file_extension == '.docx':
                        
                        try:
                            import docx
                            import io
                            
                            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
                            text_parts = []
                            
                            for paragraph in doc.paragraphs:
                                if paragraph.text and paragraph.text.strip():
                                    text_parts.append(paragraph.text.strip())
                            
                            content = "\n\n".join(text_parts)
                            
                        except Exception as docx_error:
                            print(f"DOCX extraction failed: {docx_error}")
                            content = f"Error extracting text from {uploaded_file.name}"
                    
                    else:
                        print(f"Unsupported file type: {file_extension}")
                        continue
                    
                    # Create document if we have content
                    if content and content.strip() and len(content.strip()) > 10:
                        doc = Document(
                            page_content=content.strip(),
                            metadata={"source": uploaded_file.name}
                        )
                        documents.append(doc)
                        print(f"✅ Extracted {len(content)} characters from {uploaded_file.name}")
                    else:
                        print(f"⚠️ No usable content found in {uploaded_file.name}")
                        
                except Exception as e:
                    print(f"❌ Error processing {uploaded_file.name}: {e}")
                    continue
            
            if not documents:
                print("❌ No documents could be processed")
                return False
            
            # Create vector store
            return self.create_vectorstore(documents) 
            
        except Exception as e:
            print(f"❌ Document setup failed: {e}")
            return False

    def create_vectorstore(self, documents: List[Document]) -> bool:
        """Create vector store from documents"""
        try:
            print(f"📚 Creating vector store from {len(documents)} documents...")
            
            # Split documents into chunks
            texts = self.text_splitter.split_documents(documents)
            print(f"Split into {len(texts)} chunks")
            
            if not texts:
                print("❌ No text chunks created")
                return False
            
            # Create vectorstore
            self.vectorstore = Chroma.from_documents(
                documents=texts,
                embedding=self.embeddings,
                persist_directory="./chroma_db"
            )
            
            print(f"✅ Vector store created with {len(texts)} chunks")
            return True
            
        except Exception as e:
            print(f"❌ Vector store creation failed: {e}")
            return False

    def clear_vectorstore(self):
        """Clear all documents from vector store"""
        try:
            # Clear vectorstore object
            if hasattr(self, 'vectorstore') and self.vectorstore:
                try:
                    del self.vectorstore
                except:
                    pass
            
            self.vectorstore = None
            
            # Remove database directory
            if os.path.exists("./chroma_db"):
                try:
                    shutil.rmtree("./chroma_db")
                    print("🗑️ Removed old vectorstore")
                except:
                    pass
            
            print("✅ Documents cleared")
        except Exception as e:
            print(f"Error clearing documents: {e}")

    def similarity_search(self, query: str, k: int = 4) -> List[Document]:
        """Search for similar documents"""
        if self.vectorstore is None:
            return []
        
        try:
            results = self.vectorstore.similarity_search(query, k=k)
            return results
        except Exception as e:
            print(f"Error in similarity search: {e}")
            return []

    def get_vectorstore_info(self) -> dict:
        """Get information about the current vectorstore"""
        info = {
            "vectorstore_exists": self.vectorstore is not None,
            "embeddings_model": "models/embedding-001",
            "chunk_size": self.text_splitter._chunk_size,
            "chunk_overlap": self.text_splitter._chunk_overlap
        }
        
        if self.vectorstore:
            try:                
                test_search = self.vectorstore.similarity_search("", k=1)
                info["has_documents"] = len(test_search) > 0
                info["sample_content"] = test_search[0].page_content[:100] if test_search else "No content"
            except Exception as e:
                info["error"] = str(e)
        
        return info
