from typing import Dict, List, Any, Optional
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import json
import os
import shutil
from datetime import datetime
from utils.validators import InputValidator, DateParser, TimeParser

class ConversationalForm:
    def __init__(self):
        self.reset()
    
    def reset(self):
        self.data = {
            "name": None,
            "email": None,
            "phone": None,
            "appointment_date": None,
            "appointment_time": None,
            "purpose": None
        }
        self.current_step = None
        self.validation_errors = {}
    
    def is_complete(self) -> bool:
        required_fields = ["name", "email", "phone", "appointment_date", "appointment_time"]
        return all(self.data[field] is not None for field in required_fields)
    
    def get_next_missing_field(self) -> Optional[str]:
        required_fields = ["name", "email", "phone", "appointment_date", "appointment_time", "purpose"]
        for field in required_fields:
            if self.data[field] is None:
                return field
        return None
    
    def validate_and_set_field(self, field: str, value: str) -> tuple[bool, str]:
        """Validate and set field value"""
        if field == "name":
            is_valid, result = InputValidator.validate_name(value)
            if is_valid:
                self.data["name"] = result
                return True, f"✅ Name set to: {result}"
            else:
                return False, f"❌ {result}"
        
        elif field == "email":
            is_valid, result = InputValidator.validate_email(value)
            if is_valid:
                self.data["email"] = result
                return True, f"✅ Email set to: {result}"
            else:
                return False, f"❌ {result}"
        
        elif field == "phone":
            is_valid, result = InputValidator.validate_phone(value)
            if is_valid:
                self.data["phone"] = result
                return True, f"✅ Phone set to: {result}"
            else:
                return False, f"❌ {result}"
        
        elif field == "appointment_date":
            is_valid, date_str, explanation = DateParser.parse_date_from_text(value)
            if is_valid:
                self.data["appointment_date"] = date_str
                return True, f"✅ Appointment date set to: {date_str} ({explanation})"
            else:
                return False, f"❌ {explanation}"
        
        elif field == "appointment_time":
            is_valid, time_str, explanation = TimeParser.parse_time_from_text(value)
            if is_valid:
                self.data["appointment_time"] = time_str
                return True, f"✅ Appointment time set to: {time_str} ({explanation})"
            else:
                return False, f"❌ {explanation}"
        
        elif field == "purpose":
            if len(value.strip()) > 0:
                self.data["purpose"] = value.strip()
                return True, f"✅ Purpose noted: {value.strip()}"
            else:
                return True, "✅ No specific purpose noted (optional field skipped)"
        
        return False, "❌ Unknown field"

class SimpleChatbot:
    def __init__(self, google_api_key: str):
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            google_api_key=google_api_key,
            temperature=0.1
        )
        
        # Simple document processing
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=google_api_key
        )
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,  
            chunk_overlap=50,  
            length_function=len,
        )
        
        self.vectorstore = None
        self.conversational_form = ConversationalForm()

    def setup_documents_direct(self, uploaded_files):
        """Setup documents directly from Streamlit uploaded files - NO TEMP FILES"""
        try:
            print(f"🔧 Processing {len(uploaded_files)} files directly...")
            
            # Clear old documents first
            self.clear_documents()
            
            documents = []
            
            for uploaded_file in uploaded_files:
                try:
                    print(f"📄 Processing: {uploaded_file.name}")
                    
                    # Get file content as text
                    content = ""
                    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
                    
                    if file_extension == '.txt':
                        # Text files - decode bytes to string
                        content = uploaded_file.getvalue().decode('utf-8')
                        
                    elif file_extension == '.pdf':
                        # For PDFs, simple text extraction
                        try:
                            import PyPDF2
                            import io
                            
                            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.getvalue()))
                            text_parts = []
                            
                            for page in pdf_reader.pages:
                                page_text = page.extract_text()
                                if page_text and page_text.strip():
                                    text_parts.append(page_text.strip())
                            
                            content = "\n\n".join(text_parts)
                            
                        except Exception as pdf_error:
                            print(f"PDF extraction failed: {pdf_error}")
                            content = f"Error extracting text from {uploaded_file.name}"
                    
                    elif file_extension == '.docx':
                        # For DOCX, simple text extraction
                        try:
                            import docx
                            import io
                            
                            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
                            text_parts = []
                            
                            for paragraph in doc.paragraphs:
                                if paragraph.text and paragraph.text.strip():
                                    text_parts.append(paragraph.text.strip())
                            
                            content = "\n\n".join(text_parts)
                            
                        except Exception as docx_error:
                            print(f"DOCX extraction failed: {docx_error}")
                            content = f"Error extracting text from {uploaded_file.name}"
                    
                    else:
                        print(f"Unsupported file type: {file_extension}")
                        continue
                    
                    # Create document if we have content
                    if content and content.strip() and len(content.strip()) > 10:
                        doc = Document(
                            page_content=content.strip(),
                            metadata={"source": uploaded_file.name}
                        )
                        documents.append(doc)
                        print(f"✅ Extracted {len(content)} characters from {uploaded_file.name}")
                    else:
                        print(f"⚠️ No usable content found in {uploaded_file.name}")
                        
                except Exception as e:
                    print(f"❌ Error processing {uploaded_file.name}: {e}")
                    continue
            
            if not documents:
                print("❌ No documents could be processed")
                return False
            
            # Create vector store
            print(f"📚 Creating vector store from {len(documents)} documents...")
            
            # Split documents into chunks
            texts = self.text_splitter.split_documents(documents)
            print(f"Split into {len(texts)} chunks")
            
            if not texts:
                print("❌ No text chunks created")
                return False
            
            # Create vectorstore
            self.vectorstore = Chroma.from_documents(
                documents=texts,
                embedding=self.embeddings,
                persist_directory="./chroma_db"
            )
            
            print(f"✅ Vector store created with {len(texts)} chunks")
            return True
            
        except Exception as e:
            print(f"❌ Document setup failed: {e}")
            return False

    def clear_documents(self):
        """Clear all documents from vector store"""
        try:
            # Clear vectorstore object
            if hasattr(self, 'vectorstore') and self.vectorstore:
                try:
                    del self.vectorstore
                except:
                    pass
            
            self.vectorstore = None
            
            # Remove database directory
            if os.path.exists("./chroma_db"):
                try:
                    shutil.rmtree("./chroma_db")
                    print("🗑️ Removed old vectorstore")
                except:
                    pass
            
            print("✅ Documents cleared")
        except Exception as e:
            print(f"Error clearing documents: {e}")

    def chat(self, user_input: str) -> str:
        """Main chat function"""
        user_lower = user_input.lower().strip()        
      
        # Handle booking flow first
        if self.conversational_form.current_step == "collecting":
            return self._handle_booking_flow(user_input)
        
        # Handle booking requests
        booking_keywords = ['book', 'appointment', 'call me', 'contact me', 'schedule', 'meeting']
        if any(keyword in user_lower for keyword in booking_keywords):
            self.conversational_form.current_step = "collecting"
            return "I'd be happy to help you book an appointment! 📅\n\nLet's start with your full name:"
        
        # Handle reset requests
        if any(word in user_lower for word in ["reset", "start over", "clear", "restart"]):
            self.conversational_form.reset()
            return ("🔄 I've reset everything. How can I help you today?\n\n"
                   "• Ask questions about uploaded documents\n"
                   "• Say 'I want to book an appointment' to schedule a meeting")
        
        # Search documents if available
        if self.vectorstore:
            print("📄 Searching documents...")
            try:
                return self._search_documents(user_input)
            except Exception as e:
                print(f"❌ Document search failed: {e}")
                return f"📄 Document search failed: {str(e)}"
        
        # No documents loaded
        return ("📄 No documents are currently uploaded. Please upload documents to get started!\n\n"
               "I can also help you:\n"
               "• Book appointments (say 'I want to book an appointment')")

    def _handle_booking_flow(self, user_input: str) -> str:
        """Handle booking conversation flow"""
        next_field = self.conversational_form.get_next_missing_field()
               
        if next_field is None:
            summary = self._format_booking_summary()
            self.conversational_form.current_step = "complete"
            return f"🎉 Perfect! Here's your booking information:\n\n{summary}\n\n📞 We'll contact you soon!"
        
        if next_field == "purpose" and any(word in user_input.lower() for word in ["skip", "no", "none", "nothing"]):
            self.conversational_form.data["purpose"] = "Not specified"
            summary = self._format_booking_summary()
            self.conversational_form.current_step = "complete"
            return f"✅ No problem! Here's your booking information:\n\n{summary}\n\n📞 We'll contact you soon!"
        
        # Validate and set the current field
        success, message = self.conversational_form.validate_and_set_field(next_field, user_input)
        
        if success:
            next_missing = self.conversational_form.get_next_missing_field()
            
            if next_missing:
                prompt = self._get_field_prompt(next_missing)
                return f"{message}\n\n{prompt}"
            else:
                summary = self._format_booking_summary()
                self.conversational_form.current_step = "complete"
                return f"{message}\n\n🎉 Perfect! Here's your booking:\n\n{summary}\n\n📞 We'll contact you soon!"
        else:
            prompt = self._get_field_prompt(next_field)
            return f"{message}\n\n{prompt}"

    def _search_documents(self, query: str) -> str:
        """Search through documents"""
        try:
            if not self.vectorstore:
                return "📄 No documents loaded."
            
            # Try direct search
            results = self.vectorstore.similarity_search(query, k=3)
            
            # If no results, try broader search
            if not results:
                keywords = query.lower().split()
                for keyword in keywords[:3]:
                    if len(keyword) > 3:
                        results = self.vectorstore.similarity_search(keyword, k=3)
                        if results:
                            break
            
            # If still no results, get any content
            if not results:
                results = self.vectorstore.similarity_search("", k=3)
            
            if results:
                response = "📄 **Here's what I found:**\n\n"
                
                for i, doc in enumerate(results[:2]):
                    content = doc.page_content.strip()
                    if len(content) > 300:
                        content = content[:300] + "..."
                    
                    response += f"{content}\n\n"
                
                response += "❓ Would you like me to search for something specific?"
                return response
            else:
                return "📄 I couldn't find relevant content. Try asking about specific topics in your document."
                
        except Exception as e:
            print(f"❌ Error in document search: {e}")
            return f"📄 Search error: {str(e)}"

    def _get_field_prompt(self, field: str) -> str:
        """Get prompt for each field"""
        prompts = {
            "name": "What's your full name? 👤",
            "email": "What's your email address? 📧",
            "phone": "What's your phone number? 📱",
            "appointment_date": "When would you like your appointment? 📅\n(e.g., 'tomorrow', 'next Monday', '2024-03-15')",
            "appointment_time": "What time works for you? 🕐\n(e.g., '2:30 PM', '9am', 'morning')",
            "purpose": "What's the purpose of your appointment? 🎯\n(optional - say 'skip' if you prefer not to specify)"
        }
        return prompts.get(field, "Please provide the information:")

    def _format_booking_summary(self) -> str:
        """Format booking summary"""
        data = self.conversational_form.data
        summary = f"""📋 **Appointment Summary**
                
👤 **Name:** {data['name']}
📧 **Email:** {data['email']}
📞 **Phone:** {data['phone']}
📅 **Date:** {data['appointment_date']}
🕐 **Time:** {data['appointment_time']}"""
        
        if data['purpose'] and data['purpose'] != "Not specified":
            summary += f"\n🎯 **Purpose:** {data['purpose']}"
        
        return summary
    
    def get_booking_status(self) -> Dict[str, Any]:
        """Get booking status"""
        return {
            "is_complete": self.conversational_form.is_complete(),
            "current_step": self.conversational_form.current_step,
            "data": self.conversational_form.data,
            "next_field": self.conversational_form.get_next_missing_field()
        }